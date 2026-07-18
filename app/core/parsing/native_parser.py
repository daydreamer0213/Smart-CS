"""Structured parsers for formats that do not need an advanced PDF adapter."""

import io
import re
from pathlib import Path

from app.core.parsing.contracts import ParsedDocument, ParsedElement


def _markdown_table(rows: list[list[object | None]]) -> str:
    rendered = [
        ["" if value is None else str(value).replace("|", "\\|").replace("\n", " ") for value in row]
        for row in rows
    ]
    if not rendered:
        return "| |\n| --- |"
    width = max(len(row) for row in rendered)
    rendered = [row + [""] * (width - len(row)) for row in rendered]
    lines = ["| " + " | ".join(rendered[0]) + " |"]
    lines.append("| " + " | ".join(["---"] * width) + " |")
    lines.extend("| " + " | ".join(row) + " |" for row in rendered[1:])
    return "\n".join(lines)


def _text_elements(text: str, markdown: bool) -> list[ParsedElement]:
    elements: list[ParsedElement] = []
    section_path: list[str] = []
    for block in re.split(r"\n\s*\n", text.strip()):
        block = block.strip()
        if not block:
            continue
        heading = re.fullmatch(r"(#{1,6})\s+(.+?)\s*#*", block) if markdown else None
        if heading:
            level = len(heading.group(1))
            title = heading.group(2).strip()
            section_path = section_path[: level - 1] + [title]
            elements.append(
                ParsedElement(text=title, element_type="heading", section_path=section_path.copy())
            )
        else:
            elements.append(
                ParsedElement(text=block, element_type="paragraph", section_path=section_path.copy())
            )
    return elements


def parse_text_document(filename: str, data: bytes) -> ParsedDocument:
    extension = Path(filename).suffix.lower()
    text = data.decode("utf-8", errors="replace")
    return ParsedDocument(
        parser_name="native-text",
        parser_version="1",
        page_count=0,
        elements=_text_elements(text, markdown=extension == ".md"),
    )


def _docx_body_elements(document) -> list[ParsedElement]:
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    elements: list[ParsedElement] = []
    section_path: list[str] = []
    for child in document.element.body.iterchildren():
        if child.tag.endswith("}p"):
            paragraph = Paragraph(child, document)
            text = paragraph.text.strip()
            if not text:
                continue
            style = paragraph.style.name if paragraph.style else ""
            if style.lower() == "title":
                elements.append(ParsedElement(text=text, element_type="title"))
                continue
            heading = re.fullmatch(r"Heading\s+(\d+)", style, re.IGNORECASE)
            if heading:
                level = int(heading.group(1))
                section_path = section_path[: level - 1] + [text]
                elements.append(
                    ParsedElement(text=text, element_type="heading", section_path=section_path.copy())
                )
            else:
                elements.append(
                    ParsedElement(text=text, element_type="paragraph", section_path=section_path.copy())
                )
        elif child.tag.endswith("}tbl"):
            table = Table(child, document)
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if rows:
                markdown = _markdown_table(rows)
                elements.append(
                    ParsedElement(
                        text=markdown,
                        element_type="table",
                        table_markdown=markdown,
                        section_path=section_path.copy(),
                    )
                )
    return elements


def parse_docx_document(data: bytes) -> ParsedDocument:
    from docx import Document as DocxDocument

    document = DocxDocument(io.BytesIO(data))
    return ParsedDocument(
        parser_name="python-docx-native",
        parser_version="1",
        page_count=0,
        elements=_docx_body_elements(document),
    )


def parse_xlsx_document(data: bytes) -> ParsedDocument:
    from openpyxl import load_workbook

    workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    try:
        elements = []
        for worksheet in workbook.worksheets:
            rows = list(
                worksheet.iter_rows(
                    min_row=1,
                    max_row=worksheet.max_row,
                    min_col=1,
                    max_col=worksheet.max_column,
                    values_only=True,
                )
            )
            if not any(any(cell is not None for cell in row) for row in rows):
                continue
            markdown = _markdown_table([list(row) for row in rows])
            elements.append(
                ParsedElement(
                    text=markdown,
                    element_type="table",
                    table_markdown=markdown,
                    metadata={
                        "sheet_name": worksheet.title,
                        "row_start": 1,
                        "row_end": worksheet.max_row,
                    },
                )
            )
    finally:
        workbook.close()
    return ParsedDocument(
        parser_name="openpyxl-native",
        parser_version="1",
        page_count=0,
        elements=elements,
    )


def parse_pdf_document(data: bytes) -> ParsedDocument:
    import fitz

    document = fitz.open(stream=data, filetype="pdf")
    try:
        elements = [
            ParsedElement(
                text=page.get_text().strip(),
                element_type="paragraph",
                page_start=page_number,
                page_end=page_number,
            )
            for page_number, page in enumerate(document, start=1)
        ]
        return ParsedDocument(
            parser_name="pymupdf-native",
            parser_version=fitz.VersionBind,
            page_count=document.page_count,
            elements=elements,
        )
    finally:
        document.close()


def parse_native_file(filename: str, data: bytes) -> ParsedDocument:
    extension = Path(filename).suffix.lower()
    if extension in {".txt", ".md"}:
        return parse_text_document(filename, data)
    if extension == ".docx":
        return parse_docx_document(data)
    if extension == ".xlsx":
        return parse_xlsx_document(data)
    if extension == ".pdf":
        return parse_pdf_document(data)
    raise ValueError(f"Unsupported file type: {extension}")
