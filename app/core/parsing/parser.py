"""File parsers: extract plain text from pdf, docx, xlsx, txt, md."""

import io
from pathlib import Path


def parse_pdf(data: bytes) -> str:
    """Extract text from a PDF using PyMuPDF."""
    import fitz  # PyMuPDF
    doc = fitz.open(stream=data, filetype="pdf")
    if doc.page_count == 0:
        raise ValueError("PDF has no pages")
    parts = []
    for page in doc:
        text = page.get_text().strip()
        if text:
            parts.append(text)
    doc.close()
    result = "\n\n".join(parts)
    if not result:
        raise ValueError("PDF contains no text layer (likely scanned image)")
    return result


def parse_docx(data: bytes) -> str:
    """Extract text from a Word document."""
    from docx import Document as DocxDocument
    doc = DocxDocument(io.BytesIO(data))
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Check paragraph style for heading detection
            style = para.style.name if para.style else ""
            if style.startswith("Heading") or style.startswith("heading"):
                parts.append(f"## {text}")
            else:
                parts.append(text)
    # Extract table text too
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def parse_xlsx(data: bytes) -> str:
    """Extract text from Excel — one line per row, first row as headers."""
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(data), read_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        headers = []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if all(c is None for c in row):
                continue  # skip empty rows
            if i == 0:
                headers = [str(c) if c else "" for c in row]
                continue
            cells = [str(c) if c else "" for c in row]
            # Build "Q: col1 A: col2" type lines for FAQ-style sheets
            if len(headers) >= 2 and len(cells) >= 2:
                parts.append(f"Q: {cells[0]} A: {cells[1]}")
            else:
                parts.append(" | ".join(cells))
    wb.close()
    return "\n".join(parts)


def parse_text(data: bytes) -> str:
    return data.decode("utf-8", errors="replace").strip()


_PARSERS = {
    "pdf": parse_pdf,
    "docx": parse_docx,
    "xlsx": parse_xlsx,
    "txt": parse_text,
    "md": parse_text,
}


def parse_file(filename: str, data: bytes) -> str:
    """Parse an uploaded file and return extracted text.

    Raises ValueError if file type is unsupported or parsing fails.
    """
    ext = Path(filename).suffix.lstrip(".").lower()
    if ext not in _PARSERS:
        raise ValueError(f"Unsupported file type: .{ext}")
    return _PARSERS[ext](data)


def parse_structured_file(filename: str, data: bytes):
    """Parse into structured elements without changing the legacy text API."""
    from app.core.parsing.router import parse_structured_file as _parse_structured_file

    return _parse_structured_file(filename, data)
