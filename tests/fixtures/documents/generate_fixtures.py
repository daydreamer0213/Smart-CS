from datetime import datetime, timezone
from pathlib import Path

import fitz
from docx import Document
from openpyxl import Workbook


ROOT = Path(__file__).parent
FIXED_TIME = datetime(2026, 1, 1, tzinfo=timezone.utc)


def add_text(page, point, text, size=11):
    page.insert_text(point, text, fontname="china-s", fontsize=size)


def prepare_output(name):
    path = ROOT / name
    path.unlink(missing_ok=True)
    return path


def save_pdf(name, pages):
    document = fitz.open()
    for draw_page in pages:
        page = document.new_page(width=595, height=842)
        draw_page(page)
    document.set_metadata({"title": name, "creationDate": "D:20260101000000Z"})
    document.save(prepare_output(name), garbage=4, deflate=True, no_new_id=True)
    document.close()


def build_clean_policy():
    def page_one(page):
        add_text(page, (72, 80), "北辰科技员工年假制度", 16)
        add_text(page, (72, 130), "制度版本：2026.1")

    def page_two(page):
        add_text(page, (72, 80), "年假标准")
        add_text(page, (72, 130), "工龄满十年的员工每年享有十天年假。")

    save_pdf("clean_policy.pdf", [page_one, page_two])


def build_repeated_headers():
    pages = []
    facts = ["年假申请应提前三个工作日提交。", "病假应提供医疗证明。", "婚假应一次性休完。"]
    for index, fact in enumerate(facts, start=1):
        def draw(page, page_number=index, policy_fact=fact):
            add_text(page, (72, 50), "北辰科技人力资源制度")
            add_text(page, (72, 120), policy_fact)
            add_text(page, (250, 810), f"第 {page_number} 页")

        pages.append(draw)
    save_pdf("repeated_headers.pdf", pages)


def build_leave_table():
    def draw(page):
        add_text(page, (72, 70), "年假天数对照表", 16)
        xs = [72, 250, 430]
        ys = [110, 150, 190, 230, 270]
        for x in xs:
            page.draw_line((x, ys[0]), (x, ys[-1]))
        for y in ys:
            page.draw_line((xs[0], y), (xs[-1], y))
        rows = [("工龄", "年假天数"), ("0-9年", "5天"), ("10-19年", "10天"), ("20年以上", "15天")]
        for row_index, row in enumerate(rows):
            add_text(page, (82, 137 + row_index * 40), row[0])
            add_text(page, (270, 137 + row_index * 40), row[1])

    save_pdf("leave_table.pdf", [draw])


def image_only_page(text):
    source = fitz.open()
    page = source.new_page(width=595, height=842)
    add_text(page, (72, 100), text)
    pixmap = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
    image = pixmap.tobytes("png")
    source.close()
    return image


def build_scanned_policy():
    image = image_only_page("育儿假每年五天。")
    output = fitz.open()
    image_page = output.new_page(width=595, height=842)
    image_page.insert_image(image_page.rect, stream=image)
    output.save(prepare_output("scanned_policy.pdf"), garbage=4, deflate=True, no_new_id=True)
    output.close()


def build_mixed_policy():
    output = fitz.open()
    text_page = output.new_page(width=595, height=842)
    add_text(text_page, (72, 100), "驻外员工住宿由公司统一安排。")
    image_page = output.new_page(width=595, height=842)
    image_page.insert_image(image_page.rect, stream=image_only_page("驻外补贴标准为每月3000元。"))
    output.save(prepare_output("mixed_policy.pdf"), garbage=4, deflate=True, no_new_id=True)
    output.close()


def build_two_column_policy():
    def draw(page):
        add_text(page, (72, 70), "入职办理", 15)
        add_text(page, (72, 110), "新员工应在首日完成身份核验。")
        add_text(page, (320, 70), "离职办理", 15)
        add_text(page, (320, 110), "离职员工应在三天内归还设备。")

    save_pdf("two_column_policy.pdf", [draw])


def build_encrypted_policy():
    document = fitz.open()
    page = document.new_page(width=595, height=842)
    add_text(page, (72, 100), "加密制度测试内容。")
    document.save(
        prepare_output("encrypted_policy.pdf"),
        encryption=fitz.PDF_ENCRYPT_AES_256,
        owner_pw="fixture-owner",
        user_pw="fixture-password",
        no_new_id=True,
    )
    document.close()


def build_docx():
    document = Document()
    document.core_properties.created = FIXED_TIME
    document.core_properties.modified = FIXED_TIME
    document.add_heading("员工证明开具制度", level=1)
    document.add_heading("办理时效", level=2)
    document.add_paragraph("在职证明应在两个工作日内开具。")
    document.save(prepare_output("headed_policy.docx"))


def build_xlsx():
    workbook = Workbook()
    workbook.properties.created = FIXED_TIME
    workbook.properties.modified = FIXED_TIME
    contacts = workbook.active
    contacts.title = "HR联系人"
    contacts.append(["部门", "联系人"])
    contacts.append(["薪酬福利", "payroll@example.test"])
    offices = workbook.create_sheet("办公地点")
    offices.append(["城市", "HRBP"])
    offices.append(["上海", "测试联系人"])
    workbook.save(prepare_output("hr_contacts.xlsx"))


def main():
    build_clean_policy()
    build_repeated_headers()
    build_leave_table()
    build_scanned_policy()
    build_mixed_policy()
    build_two_column_policy()
    build_encrypted_policy()
    build_docx()
    build_xlsx()


if __name__ == "__main__":
    main()
