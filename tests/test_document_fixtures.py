import json
from pathlib import Path

import fitz
import pytest
from docx import Document
from openpyxl import load_workbook


FIXTURE_DIR = Path(__file__).parent / "fixtures" / "documents"
REQUIRED_ENTRY_FIELDS = {
    "id",
    "filename",
    "format",
    "category",
    "expected_baseline_status",
    "expected_chunk_count",
    "expected_missing_facts",
    "required_facts",
}


def load_manifest():
    return json.loads((FIXTURE_DIR / "manifest.json").read_text(encoding="utf-8"))


def assert_manifest_entry_structure(entries):
    assert len(entries) == 9
    assert len({entry["id"] for entry in entries}) == len(entries)
    assert len({entry["filename"] for entry in entries}) == len(entries)
    assert all(REQUIRED_ENTRY_FIELDS <= entry.keys() for entry in entries)
    assert all("expected_page_count" in entry for entry in entries if entry["format"] == "pdf")


def test_document_fixture_manifest_covers_enterprise_shapes():
    manifest = load_manifest()
    entries = manifest["fixtures"]

    assert manifest["schema_version"] == 1
    assert {entry["category"] for entry in entries} == {
        "clean_text",
        "repeated_header_footer",
        "table",
        "scanned",
        "mixed_text_scan",
        "two_column",
        "encrypted",
        "headed_docx",
        "multi_sheet_xlsx",
    }
    assert all((FIXTURE_DIR / entry["filename"]).is_file() for entry in entries)
    assert all(entry["required_facts"] for entry in entries if entry["category"] != "encrypted")


def test_document_fixture_manifest_entries_are_unique_complete_and_match_pdf_pages():
    entries = load_manifest()["fixtures"]

    assert_manifest_entry_structure(entries)
    for entry in entries:
        if entry["format"] == "pdf":
            document = fitz.open(FIXTURE_DIR / entry["filename"])
            assert document.page_count == entry["expected_page_count"]
            document.close()


def test_manifest_entry_structure_rejects_duplicate_ids():
    entries = [dict(entry) for entry in load_manifest()["fixtures"]]
    entries[1]["id"] = entries[0]["id"]

    with pytest.raises(AssertionError):
        assert_manifest_entry_structure(entries)


def test_text_pdf_fixtures_expose_all_required_facts():
    entries = {entry["category"]: entry for entry in load_manifest()["fixtures"]}

    for category in ("clean_text", "repeated_header_footer", "table", "two_column"):
        entry = entries[category]
        document = fitz.open(FIXTURE_DIR / entry["filename"])
        text = "".join(page.get_text() for page in document)
        assert all(fact in text for fact in entry["required_facts"])
        document.close()


def test_repeated_header_and_numbered_footer_are_positioned_on_every_page():
    entry = next(
        entry
        for entry in load_manifest()["fixtures"]
        if entry["category"] == "repeated_header_footer"
    )
    document = fitz.open(FIXTURE_DIR / entry["filename"])

    for page_number, page in enumerate(document, start=1):
        headers = page.search_for("北辰科技人力资源制度")
        footers = page.search_for(f"第 {page_number} 页")
        assert len(headers) == 1
        assert len(footers) == 1
        assert headers[0].y1 < page.rect.height * 0.1
        assert footers[0].y0 > page.rect.height * 0.9

    document.close()


def test_leave_table_contains_all_rows_inside_a_drawn_grid():
    entry = next(
        entry for entry in load_manifest()["fixtures"] if entry["category"] == "table"
    )
    document = fitz.open(FIXTURE_DIR / entry["filename"])
    page = document[0]
    rows = [
        ("工龄", "年假天数"),
        ("0-9年", "5天"),
        ("10-19年", "10天"),
        ("20年以上", "15天"),
    ]
    word_rects = {
        word[4]: fitz.Rect(word[:4])
        for word in page.get_text("words")
        if 110 < word[1] < 270
    }

    row_rects = []
    for left_text, right_text in rows:
        left = word_rects[left_text]
        right = word_rects[right_text]
        assert left.x1 < right.x0
        assert left.y0 == pytest.approx(right.y0, abs=1)
        row_rects.append((left, right))

    assert [left.y0 for left, _ in row_rects] == sorted(
        left.y0 for left, _ in row_rects
    )

    lines = [
        item[1:]
        for drawing in page.get_drawings()
        for item in drawing["items"]
        if item[0] == "l"
    ]
    vertical = sorted(
        (start.x, start.y, end.y)
        for start, end in lines
        if start.x == pytest.approx(end.x)
    )
    horizontal = sorted(
        (start.y, start.x, end.x)
        for start, end in lines
        if start.y == pytest.approx(end.y)
    )
    assert vertical == pytest.approx(
        [(72, 110, 270), (250, 110, 270), (430, 110, 270)]
    )
    assert horizontal == pytest.approx(
        [
            (110, 72, 430),
            (150, 72, 430),
            (190, 72, 430),
            (230, 72, 430),
            (270, 72, 430),
        ]
    )
    assert all(72 < left.x0 < 250 < right.x0 < 430 for left, right in row_rects)
    assert all(110 < left.y0 < 270 for left, _ in row_rects)
    document.close()


def test_two_column_facts_are_drawn_on_opposite_page_halves():
    entry = next(
        entry
        for entry in load_manifest()["fixtures"]
        if entry["category"] == "two_column"
    )
    document = fitz.open(FIXTURE_DIR / entry["filename"])
    page = document[0]
    left = page.search_for("新员工应在首日完成身份核验。")
    right = page.search_for("离职员工应在三天内归还设备。")

    assert len(left) == len(right) == 1
    assert left[0].x1 < page.rect.width / 2
    assert right[0].x0 > page.rect.width / 2
    assert left[0].y0 == pytest.approx(right[0].y0, abs=1)
    document.close()


def test_scanned_and_mixed_pdf_page_layers_match_their_baselines():
    entries = {entry["category"]: entry for entry in load_manifest()["fixtures"]}

    scanned = fitz.open(FIXTURE_DIR / entries["scanned"]["filename"])
    assert not scanned[0].get_text().strip()
    assert scanned[0].get_images(full=True)
    scanned.close()

    mixed_entry = entries["mixed_text_scan"]
    mixed = fitz.open(FIXTURE_DIR / mixed_entry["filename"])
    assert mixed_entry["required_facts"][0] in mixed[0].get_text()
    assert not mixed[1].get_text().strip()
    assert mixed[1].get_images(full=True)
    mixed.close()


def test_encrypted_pdf_requires_the_fixture_password():
    entry = next(entry for entry in load_manifest()["fixtures"] if entry["category"] == "encrypted")
    document = fitz.open(FIXTURE_DIR / entry["filename"])

    assert document.needs_pass
    assert not document.authenticate("wrong-password")
    assert document.authenticate("fixture-password")
    document.close()


def test_docx_and_xlsx_fixtures_preserve_declared_structure_and_facts():
    entries = {entry["category"]: entry for entry in load_manifest()["fixtures"]}

    docx_entry = entries["headed_docx"]
    document = Document(FIXTURE_DIR / docx_entry["filename"])
    assert {"Heading 1", "Heading 2"} <= {paragraph.style.name for paragraph in document.paragraphs}
    docx_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert all(fact in docx_text for fact in docx_entry["required_facts"])

    xlsx_entry = entries["multi_sheet_xlsx"]
    workbook = load_workbook(FIXTURE_DIR / xlsx_entry["filename"], read_only=True, data_only=True)
    assert workbook.sheetnames == ["HR联系人", "办公地点"]
    values = [
        str(cell.value)
        for sheet in workbook.worksheets
        for row in sheet.iter_rows()
        for cell in row
        if cell.value is not None
    ]
    assert all(fact in values for fact in xlsx_entry["required_facts"])
    workbook.close()
