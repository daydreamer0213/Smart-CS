import io

from docx import Document
from openpyxl import Workbook


def test_structured_text_and_markdown_keep_paragraphs_and_section_paths():
    from app.core.parsing.router import parse_structured_file

    text = parse_structured_file("policy.txt", b"First paragraph.\n\nSecond paragraph.")
    markdown = parse_structured_file(
        "policy.md", b"# Handbook\n\nIntroduction.\n\n## Leave\n\nTen days.",
    )

    assert text.page_count == 0
    assert [element.element_type for element in text.elements] == ["paragraph", "paragraph"]
    assert [element.text for element in text.elements] == ["First paragraph.", "Second paragraph."]
    assert [(element.element_type, element.section_path) for element in markdown.elements] == [
        ("heading", ["Handbook"]),
        ("paragraph", ["Handbook"]),
        ("heading", ["Handbook", "Leave"]),
        ("paragraph", ["Handbook", "Leave"]),
    ]


def test_structured_docx_preserves_body_order_headings_and_table_markdown():
    from app.core.parsing.router import parse_structured_file

    source = Document()
    source.add_heading("Employee handbook", level=0)
    source.add_heading("Handbook", level=1)
    source.add_paragraph("Opening paragraph.")
    table = source.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Years"
    table.cell(0, 1).text = "Days"
    table.cell(1, 0).text = "10"
    table.cell(1, 1).text = "10"
    source.add_heading("Leave", level=2)
    source.add_paragraph("Closing paragraph.")
    data = io.BytesIO()
    source.save(data)

    document = parse_structured_file("handbook.docx", data.getvalue())

    assert [element.element_type for element in document.elements] == [
        "title", "heading", "paragraph", "table", "heading", "paragraph",
    ]
    assert document.elements[3].table_markdown == "| Years | Days |\n| --- | --- |\n| 10 | 10 |"
    assert document.elements[3].section_path == ["Handbook"]
    assert document.elements[5].section_path == ["Handbook", "Leave"]


def test_structured_xlsx_keeps_each_non_empty_sheet_full_table_and_row_bounds():
    from app.core.parsing.router import parse_structured_file

    workbook = Workbook()
    first = workbook.active
    first.title = "Contacts"
    first.append(["Department", "Email", "Region"])
    first.append(["Benefits", "benefits@example.test", "East"])
    second = workbook.create_sheet("Offices")
    second.append(["City", "Owner", "Capacity"])
    second.append(["Shanghai", "HRBP", 20])
    data = io.BytesIO()
    workbook.save(data)

    document = parse_structured_file("contacts.xlsx", data.getvalue())

    assert [element.metadata for element in document.elements] == [
        {"sheet_name": "Contacts", "row_start": 1, "row_end": 2},
        {"sheet_name": "Offices", "row_start": 1, "row_end": 2},
    ]
    assert all(element.element_type == "table" for element in document.elements)
    assert "| Department | Email | Region |" in document.elements[0].table_markdown
    assert "| Shanghai | HRBP | 20 |" in document.elements[1].table_markdown
